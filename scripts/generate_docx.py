import sys
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

import urllib.request
import base64
import json

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_horizontal_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12') # 1.5 pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_mermaid_diagram(doc, mermaid_code):
    try:
        json_obj = {"code": mermaid_code, "mermaid": {"theme": "default"}}
        json_str = json.dumps(json_obj)
        b64_bytes = base64.urlsafe_b64encode(json_str.encode('utf-8'))
        b64_str = b64_bytes.decode('utf-8').replace('=', '')
        
        url = f"https://mermaid.ink/img/{b64_str}"
        
        req = urllib.request.Request(
            url, 
            headers={'User-Agent': 'Mozilla/5.0'}
        )
        with urllib.request.urlopen(req, timeout=15) as response:
            img_data = response.read()
            
        temp_img_path = "temp_mermaid_scope.png"
        with open(temp_img_path, 'wb') as f:
            f.write(img_data)
            
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(temp_img_path, width=Inches(6.0))
        
        os.remove(temp_img_path)
        print("Successfully rendered and inserted Mermaid diagram into document.")
        return True
    except Exception as e:
        print(f"Warning: Failed to render Mermaid diagram via mermaid.ink: {e}")
        # If it fails, we fall back to printing the raw code as a block
        p = doc.add_paragraph()
        run = p.add_run(f"[Mermaid Diagram]\n{mermaid_code}")
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        return False

def build_document():
    doc = Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    styles = doc.styles
    
    # Configure Normal Style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)
    
    # Document Title (Heading 1 equivalent for doc title)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(18)
    title_run = title_p.add_run("Project Scope & Internship Roadmap 2026: Mathchines")
    title_run.font.name = 'Calibri Light'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy color
    
    # Subtitle or metadata
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(24)
    run_meta = meta_p.add_run("Prepared for Desmond Appiah\nBuild with AI Internship Program 2026\nDate: July 15, 2026")
    run_meta.font.italic = True
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    add_horizontal_border(meta_p)
    
    # Read Markdown file
    md_path = "PROJECT_SCOPE.md"
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        sys.exit(1)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_headers = []
    table_rows = []
    
    # Skip Title and horizontal line if we match them in md
    skip_heading_1 = True # We already printed custom styled title
    in_mermaid = False
    mermaid_lines = []
    
    for i, line in enumerate(lines):
        clean_line = line.strip()
        
        # Check for mermaid diagrams and render them
        if clean_line.startswith('```mermaid'):
            in_mermaid = True
            mermaid_lines = []
            continue
        elif in_mermaid and clean_line == '```':
            in_mermaid = False
            mermaid_code = "\n".join(mermaid_lines)
            add_mermaid_diagram(doc, mermaid_code)
            continue
        elif in_mermaid:
            mermaid_lines.append(line)
            continue
        
        # Check for tables
        if clean_line.startswith('|') and clean_line.endswith('|'):
            if '---' in clean_line:
                # Separator line, skip
                continue
            cells = [c.strip() for c in clean_line.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            continue
        elif in_table:
            # End of table, process it
            in_table = False
            # Render Table
            doc_table = doc.add_table(rows=len(table_rows) + 1, cols=len(table_headers))
            doc_table.style = 'Light Shading Accent 1'
            doc_table.autofit = True
            
            # Format header
            hdr_cells = doc_table.rows[0].cells
            for idx, text in enumerate(table_headers):
                hdr_cells[idx].text = text
                set_cell_background(hdr_cells[idx], "1B365D") # Navy header
                set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
                # Bold text and white color
                for p in hdr_cells[idx].paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.name = 'Calibri'
                        r.font.size = Pt(10)
                        
            # Format rows
            for r_idx, row in enumerate(table_rows):
                row_cells = doc_table.rows[r_idx + 1].cells
                fill_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF" # Zebra striping
                for c_idx, text in enumerate(row):
                    row_cells[c_idx].text = text
                    set_cell_background(row_cells[c_idx], fill_color)
                    set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=150, right=150)
                    for p in row_cells[c_idx].paragraphs:
                        p.paragraph_format.space_after = Pt(0)
                        for r in p.runs:
                            r.font.name = 'Calibri'
                            r.font.size = Pt(9.5)
                            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                            
            # Add an empty paragraph after table
            doc.add_paragraph()
            table_headers = []
            table_rows = []
            
        # Parse standard lines
        if clean_line.startswith('# '):
            text = clean_line[2:]
            if skip_heading_1:
                # We skip the very first Heading 1 since it's the title we custom-added
                skip_heading_1 = False
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.font.name = 'Calibri Light'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            add_horizontal_border(p)
            
        elif clean_line.startswith('## '):
            text = clean_line[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.font.name = 'Calibri Light'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x5B, 0x88)
            
        elif clean_line.startswith('### '):
            text = clean_line[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            
        elif clean_line.startswith('* ') or clean_line.startswith('- '):
            # List item
            text = clean_line[2:]
            # Remove markdown links e.g. [text](url) -> text
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            # Find bold elements inside
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
                    
        elif clean_line.startswith('1. ') or clean_line.startswith('2. ') or re.match(r'^\d+\.\s', clean_line):
            # Ordered List item
            text = re.sub(r'^\d+\.\s+', '', clean_line)
            text = re.sub(r'^\d+\.\s+', '', text)
            # Remove markdown links
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
                    
        elif clean_line == '---':
            # Horizontal rule
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            add_horizontal_border(p)
            
        elif clean_line == '' or clean_line.startswith('```'):
            # Empty or code block boundary, skip or add minor space if appropriate
            continue
            
        else:
            # Paragraph
            # Remove markdown links
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_line)
            # Handle inline bolding
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
                    
    # Save document
    output_filename = "PROJECT_SCOPE.docx"
    doc.save(output_filename)
    print(f"Successfully generated {output_filename}!")

if __name__ == '__main__':
    build_document()
